"""A beginner-friendly Word document editor using text commands.

Install the dependency first:
    pip uninstall docx
    pip install python-docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt


class WordDocumentAgent:
    """Observe a text command, plan an action, and edit a Word document."""

    def __init__(self):
        self.document = None
        self.current_file = None
        self.modified = False

    def observe(self):
        """OBSERVE: receive the user's text command."""
        self.show_menu()
        return input("\nEnter a command: ").strip()

    def plan(self, command):
        """PLAN: convert a text command into an action."""
        command = command.lower().strip()

        actions = {
            "create": "create_new",
            "create new": "create_new",
            "new": "create_new",
            "new document": "create_new",

            "open": "open_document",
            "open document": "open_document",

            "write": "write_text",
            "write text": "write_text",
            "add text": "write_text",

            "font": "change_font",
            "change font": "change_font",
            "format": "change_font",

            "save": "save_document",
            "save document": "save_document",

            "delete": "delete_document",
            "delete document": "delete_document",

            "help": "help",
            "show commands": "help",

            "exit": "exit",
            "quit": "exit",
            "stop": "exit",
        }

        return actions.get(command, "unknown")

    def act(self, action):
        """ACT: perform the selected document operation."""
        methods = {
            "create_new": self.create_new,
            "open_document": self.open_document,
            "write_text": self.write_text,
            "change_font": self.change_font,
            "save_document": self.save_document,
            "delete_document": self.delete_document,
        }

        if action == "exit":
            return self.exit_program()

        if action == "help":
            self.show_help()
            return True

        if action == "unknown":
            print("\nI do not understand that command.")
            print("Type 'help' to see the available commands.")
            return True

        methods[action]()
        return True

    def create_new(self):
        """Create a blank Word document."""
        if not self.confirm_discard_changes():
            return

        self.document = Document()
        self.current_file = None
        self.modified = False

        print("\nA new blank Word document has been created.")

    def open_document(self):
        """Open an existing Word document."""
        if not self.confirm_discard_changes():
            return

        filename = input("Enter the Word filename to open: ").strip()

        if not filename:
            print("No filename was supplied.")
            return

        path = self.get_docx_path(filename)

        if not path.is_file():
            print(f"File not found: {path}")
            return

        try:
            self.document = Document(path)
            self.current_file = path
            self.modified = False

            print(f"Opened: {path}")
        except Exception as error:
            print(f"Could not open the document: {error}")

    def write_text(self):
        """Add paragraphs to the current document."""
        if not self.require_document():
            return

        print("\nEnter your text.")
        print("Press Enter on an empty line when you have finished.")

        lines = []

        while True:
            line = input()

            if line == "":
                break

            lines.append(line)

        if not lines:
            print("No text was added.")
            return

        for line in lines:
            self.document.add_paragraph(line)

        self.modified = True
        print(f"Added {len(lines)} paragraph(s).")

    def change_font(self):
        """Change the font of existing and future document text."""
        if not self.require_document():
            return

        font_name = input(
            "Enter the font name, such as Times New Roman: "
        ).strip()

        if not font_name:
            print("The font name cannot be empty.")
            return

        size_text = input(
            "Enter the font size in points, such as fourteen or 14: "
        ).strip()

        font_size = self.parse_font_size(size_text)

        if font_size is None:
            print("Please enter a valid positive font size.")
            return

        # Set the default font for text added later.
        normal_font = self.document.styles["Normal"].font
        normal_font.name = font_name
        normal_font.size = Pt(font_size)

        # Change existing text in paragraphs and tables.
        for paragraph in self.iter_paragraphs():
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)

        self.modified = True

        print(f"Font changed to {font_name}, {font_size:g} point.")

    def save_document(self):
        """Save the current Word document."""
        if not self.require_document():
            return

        if self.current_file:
            prompt = (
                f"Enter a filename, or press Enter to save as "
                f"'{self.current_file}': "
            )
        else:
            prompt = "Enter a filename for the document: "

        entered_name = input(prompt).strip()

        if entered_name:
            path = self.get_docx_path(entered_name)
        elif self.current_file:
            path = self.current_file
        else:
            print("A filename is required for a new document.")
            return

        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            self.document.save(path)

            self.current_file = path
            self.modified = False

            print(f"Saved: {path}")
        except Exception as error:
            print(f"Could not save the document: {error}")

    def delete_document(self):
        """Delete a Word document after confirmation."""
        if self.current_file:
            prompt = (
                f"Enter the filename to delete, or press Enter for "
                f"'{self.current_file}': "
            )
        else:
            prompt = "Enter the Word filename to delete: "

        entered_name = input(prompt).strip()

        if entered_name:
            path = self.get_docx_path(entered_name)
        elif self.current_file:
            path = self.current_file
        else:
            print("No filename was supplied.")
            return

        if not path.is_file():
            print(f"File not found: {path}")
            return

        confirmation = input(
            f"Type DELETE to permanently delete '{path}': "
        ).strip()

        if confirmation != "DELETE":
            print("Delete cancelled.")
            return

        try:
            path.unlink()

            if (
                self.current_file
                and path.resolve() == self.current_file.resolve()
            ):
                self.document = None
                self.current_file = None
                self.modified = False

            print(f"Deleted: {path}")
        except Exception as error:
            print(f"Could not delete the document: {error}")

    def iter_paragraphs(self):
        """Yield paragraphs in the document body and tables."""
        yield from self.document.paragraphs

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    def require_document(self):
        """Check that a document is currently available."""
        if self.document is None:
            print(
                "Create a new document or open an existing "
                "document first."
            )
            return False

        return True

    def confirm_discard_changes(self):
        """Ask before discarding unsaved changes."""
        if self.document is not None and self.modified:
            answer = input(
                "Unsaved changes will be lost. "
                "Type YES to continue: "
            ).strip().lower()

            return answer == "yes"

        return True

    def exit_program(self):
        """Exit the application safely."""
        if not self.confirm_discard_changes():
            print("Exit cancelled.")
            return True

        print("Program over.")
        return False

    @staticmethod
    def parse_font_size(value):
        """Convert a number or a written number into a font size."""
        written_sizes = {
            "eight": 8,
            "nine": 9,
            "ten": 10,
            "eleven": 11,
            "twelve": 12,
            "thirteen": 13,
            "fourteen": 14,
            "fifteen": 15,
            "sixteen": 16,
            "eighteen": 18,
            "twenty": 20,
            "twenty two": 22,
            "twenty four": 24,
            "twenty eight": 28,
            "thirty two": 32,
            "thirty six": 36,
            "forty": 40,
            "forty eight": 48,
            "seventy two": 72,
        }

        cleaned_value = value.lower().strip().replace("-", " ")

        if cleaned_value in written_sizes:
            return written_sizes[cleaned_value]

        try:
            size = float(cleaned_value)

            if size > 0:
                return size
        except ValueError:
            pass

        return None

    @staticmethod
    def get_docx_path(name):
        """Return a Path with a .docx extension."""
        path = Path(name.strip()).expanduser()

        if path.suffix.lower() != ".docx":
            path = path.with_suffix(".docx")

        return path

    def show_menu(self):
        """Display the available text commands."""
        current = (
            str(self.current_file)
            if self.current_file
            else "No file selected"
        )

        changed = " (unsaved changes)" if self.modified else ""

        print("\n" + "=" * 55)
        print("WORD DOCUMENT AGENT")
        print(f"Current document: {current}{changed}")
        print("=" * 55)
        print("CREATE  - Create a new document")
        print("OPEN    - Open an existing document")
        print("WRITE   - Add text to the document")
        print("FONT    - Change the font name and size")
        print("SAVE    - Save the document")
        print("DELETE  - Delete a document")
        print("HELP    - Show command details")
        print("EXIT    - Close the program")

    @staticmethod
    def show_help():
        """Display command examples."""
        print("\nAVAILABLE COMMANDS")
        print("- Type 'create' to create a new document.")
        print("- Type 'open' to open a document.")
        print("- Type 'write' to add text.")
        print("- Type 'font' to change the font.")
        print("- Type 'save' to save the document.")
        print("- Type 'delete' to delete a document.")
        print("- Type 'exit' to close the program.")

    def run(self):
        """Run the Observe → Plan → Act loop."""
        running = True

        while running:
            goal = self.observe()
            print("\nObserved command:", goal)

            action = self.plan(goal)
            print("Planned action:", action)

            running = self.act(action)


if __name__ == "__main__":
    WordDocumentAgent().run()
