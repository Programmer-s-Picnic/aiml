"""WordAgent: performs delegated .docx jobs with python-docx."""

import shlex
from pathlib import Path

from docx import Document
from docx.shared import Pt


class WordAgent:
    """Specialist agent responsible only for Word-document operations."""

    def __init__(self):
        self.base_folder = Path(__file__).resolve().parent
        self.document = None
        self.current_file = None
        self.modified = False

    def Observe(self, transferred_request):
        """OBSERVE: receive a Word job transferred by HelloAgent."""
        if isinstance(transferred_request, tuple):
            if not transferred_request or not isinstance(transferred_request[0], str):
                raise TypeError("A Word command tuple must start with a command word")
            request = transferred_request
        else:
            request = transferred_request.strip()
        print(f"WordAgent observed: {request}")
        return request

    def Plan(self, request):
        """PLAN: convert the transferred request into an action and arguments."""
        if isinstance(request, tuple):
            parts = list(request)
        else:
            try:
                parts = shlex.split(request)
            except ValueError as error:
                return "command_error", [str(error)]

        if not parts:
            return "unknown", []

        command = parts[0].strip().lower()
        arguments = parts[1:]
        known_commands = {
            "create", "open", "read", "write", "font", "save", "delete", "status"
        }
        action = command if command in known_commands else "unknown"
        print(f"WordAgent planned: {action}")
        return action, arguments

    def Act(self, action, arguments):
        """ACT: perform the planned Word operation and return its result."""
        methods = {
            "create": self.create,
            "open": self.open_document,
            "read": self.read_document,
            "write": self.write,
            "font": self.change_font,
            "save": self.save,
            "delete": self.delete,
            "status": self.status,
        }
        if action == "command_error":
            return f"Command error: {arguments[0]}"
        if action == "unknown":
            return "Unknown Word command."
        try:
            return methods[action](arguments)
        except Exception as error:
            return f"Job failed: {error}"

    def execute(self, transferred_request):
        """Compatibility helper that uses the public Observe -> Plan -> Act cycle."""
        request = self.Observe(transferred_request)
        action, arguments = self.Plan(request)
        result = self.Act(action, arguments)
        print(f"WordAgent acted: {result}")
        return result

    def execute_many(self, commands):
        """Execute a tuple of commands and return one result per command."""
        if not isinstance(commands, tuple):
            return ("Batch rejected: commands must be supplied as a tuple.",)

        results = []
        for command in commands:
            if not isinstance(command, str) or not command.strip():
                results.append(f"{command!r} -> Invalid command")
                continue
            result = self.execute(command)
            results.append(f"{command} -> {result}")
        return tuple(results)

    def create(self, arguments):
        if not arguments:
            return "Use: create filename.docx"
        if self.modified:
            return "Save the current document before creating another one."

        self.document = Document()
        self.current_file = self.safe_path(arguments[0])
        self.modified = True
        return f"Created new document in memory: {self.current_file.name}"

    def open_document(self, arguments):
        if not arguments:
            return "Use: open filename.docx"
        if self.modified:
            return "Save the current document before opening another one."

        path = self.safe_path(arguments[0])
        if not path.is_file():
            return f"File not found: {path.name}"
        self.document = Document(path)
        self.current_file = path
        self.modified = False
        return f"Opened: {path.name}"

    def read_document(self, arguments):
        """Read and return all text from a saved Word document."""
        if arguments:
            path = self.safe_path(arguments[0])
        elif self.current_file:
            path = self.current_file
        else:
            return "Use: read filename.docx"

        if not path.is_file():
            return f"File not found: {path.name}"

        document = Document(path)
        text_parts = [p.text for p in document.paragraphs if p.text]
        for table in document.tables:
            for row in table.rows:
                text_parts.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(part for part in text_parts if part)
        return text if text else f"{path.name} is empty."

    def write(self, arguments):
        problem = self.require_document()
        if problem:
            return problem
        if not arguments:
            return 'Use: write "Your text here"'

        text = " ".join(str(argument) for argument in arguments)
        self.document.add_paragraph(text)
        self.modified = True
        return f"Added text: {text}"

    def change_font(self, arguments):
        problem = self.require_document()
        if problem:
            return problem
        if len(arguments) < 2:
            return 'Use: font "Times New Roman" 14'

        try:
            size = float(arguments[-1])
            if size <= 0:
                raise ValueError
        except ValueError:
            return "Font size must be a positive number."

        font_name = " ".join(str(argument) for argument in arguments[:-1])
        normal = self.document.styles["Normal"].font
        normal.name = font_name
        normal.size = Pt(size)
        for paragraph in self.iter_paragraphs():
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(size)
        self.modified = True
        return f"Changed font to {font_name}, {size:g} pt."

    def save(self, arguments):
        problem = self.require_document()
        if problem:
            return problem
        if arguments:
            self.current_file = self.safe_path(arguments[0])
        if self.current_file is None:
            return "Use: save filename.docx"

        self.document.save(self.current_file)
        self.modified = False
        return f"Saved: {self.current_file.name}"

    def delete(self, arguments):
        if arguments:
            path = self.safe_path(arguments[0])
        elif self.current_file:
            path = self.current_file
        else:
            return "Use: delete filename.docx"

        if not path.is_file():
            return f"File not found: {path.name}"
        confirmation = input(f"WordAgent: Type DELETE to delete {path.name}: ")
        if confirmation.strip() != "DELETE":
            return "Delete cancelled."

        path.unlink()
        if self.current_file == path:
            self.document = None
            self.current_file = None
            self.modified = False
        return f"Deleted: {path.name}"

    def status(self, _arguments):
        if self.document is None:
            return "No document is open."
        state = "unsaved changes" if self.modified else "saved"
        return f"Current document: {self.current_file.name} ({state})"

    def iter_paragraphs(self):
        yield from self.document.paragraphs
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    def require_document(self):
        if self.document is None:
            return "Create or open a document first."
        return None

    def has_unsaved_changes(self):
        return self.document is not None and self.modified

    def safe_path(self, filename):
        """Keep every managed .docx beside the two agents."""
        clean_name = Path(filename).name
        path = Path(clean_name)
        if path.suffix.lower() != ".docx":
            path = path.with_suffix(".docx")
        return self.base_folder / path
